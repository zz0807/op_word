from docx import Document
import csv
import os
import re


def generate_word(model_filename,  medicine):
    document = Document("模板/" + model_filename)

    filename = 'data.csv'
    sections = document.sections
    for section in sections:
        header = section.header
        for paragraph in header.paragraphs:
            with open(filename) as f:
                reader = csv.reader(f)
                for row in reader:
                    if row[1] in paragraph.text:
                        inline = paragraph.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if row[1] in inline[i].text:
                                inline[i].text = inline[i].text.replace(row[1], row[2])
                            else:
                                if i + 1 < len(inline):
                                    if row[1][1:-1] in inline[i].text and inline[i - 1].text == row[1][0] \
                                            and inline[i + 1].text == row[1][-1]:
                                        inline[i - 1].text = inline[i-1].text.replace(row[1][0], '')
                                        inline[i + 1].text = inline[i + 1].text.replace(row[1][-1], '')
                                        text = re.sub(row[1][1:-1], row[2], inline[i].text)
                                        inline[i].text = text

    for paragraph in document.paragraphs:
        with open(filename) as f:
            reader = csv.reader(f)
            for row in reader:
                if row[1] in paragraph.text:
                    inline = paragraph.runs
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        print(inline[i].text)
                        if row[1] in inline[i].text:
                            inline[i].text = inline[i].text.replace(row[1], row[2])
                        else:
                            if i+1 < len(inline):
                                if row[1][1:-1] in inline[i].text and inline[i-1].text == row[1][0] \
                                        and inline[i+1].text == row[1][-1]:
                                    inline[i - 1].text = inline[i - 1].text.replace(row[1][0], '')
                                    inline[i + 1].text = inline[i + 1].text.replace(row[1][-1], '')
                                    text = re.sub(row[1][1:-1], row[2], inline[i].text)
                                    inline[i].text = text

    for table in document.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                with open(filename) as f:
                    reader = csv.reader(f)
                    for row in reader:
                        if row[1] in paragraph.text:
                            inline = paragraph.runs
                            # Loop added to work with runs (strings with same style)
                            for i in range(len(inline)):
                                if row[1] in inline[i].text:
                                    inline[i].text = inline[i].text.replace(row[1], row[2])
                                else:
                                    if i + 1 < len(inline):
                                        if row[1][1:-1] in inline[i].text and inline[i - 1].text == row[1][0] \
                                                and inline[i + 1].text == row[1][-1]:
                                            inline[i - 1].text = inline[i - 1].text.replace(row[1][0], '')
                                            inline[i + 1].text = inline[i + 1].text.replace(row[1][-1], '')
                                            text = re.sub(row[1][1:-1], row[2], inline[i].text)
                                            inline[i].text = text

    des_name = model_filename.replace("模板", medicine)
    document.save("替换后文档/" + des_name)


print("请输入药品名称：")
medicine = input()
files = [x for x in os.listdir('模板')]
for single_file in files:
    generate_word(single_file, medicine)