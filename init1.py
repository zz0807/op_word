import re
from docx import Document
import csv

document = Document("word1.docx")
paragraphs = document.paragraphs


filename = 'data.csv'


for paragraph in document.paragraphs:
    with open(filename) as f:
        reader = csv.reader(f)
        for row in reader:
            text = re.sub(row[1], row[2], paragraph.text)
            paragraph.text = text

for table in document.tables:
    for cell in table._cells:
        for paragraph in cell.paragraphs:
            with open(filename) as f:
                reader = csv.reader(f)
                for row in reader:
                    text = re.sub(row[1], row[2], paragraph.text)
                    paragraph.text = text

document.save("word2.docx")
